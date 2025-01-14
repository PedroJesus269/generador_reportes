import streamlit as st
from PIL import Image


from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches
from docx.shared import Pt
from docx import Document
from io import BytesIO


from datetime import datetime
import matplotlib.dates as mdates
import matplotlib.patches as mpatches
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt


meses_es = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
}


# ******************************* CÓDIGO ORIGINAL DE PEDRO JESÚS
def plot_eventos(df):
    # Conversión de fechas
    df['Start'] = pd.to_datetime(df['Start'])
    df['hora'] = df['Start'].dt.hour

    # Filtrar solo las horas entre las 7:00 y las 7:00 del día siguiente
    df_filtrado = df[(df['hora'] >= 7) & (df['hora'] < 7 + 24)]

    # Contar los eventos por hora y por tipo (en el dataframe filtrado)
    contador_tipo_1 = df_filtrado[df_filtrado['Description'] == 'Caution Flash'].groupby('hora').size()  # Amarillo
    contador_tipo_2 = df_filtrado[df_filtrado['Description'] == 'Alarm Flash'].groupby('hora').size()  # Roja
    contador_tipo_3 = df_filtrado[df_filtrado['Description'] == 'Warning Flash'].groupby('hora').size()  # Naranja

    # Sumar Alarm Flash y Caution Flash
    contador_tipo_2_y_3 = contador_tipo_2.add(contador_tipo_3, fill_value=0)

    # Crear un DataFrame con ambos conteos (ajustado al rango de horas de 7 a 7)
    conteos = pd.DataFrame({
        'Amarilla': contador_tipo_1.reindex(range(7, 7 + 24), fill_value=0),
        'Roja': contador_tipo_2_y_3.reindex(range(7, 7 + 24), fill_value=0)
    }).fillna(0)

    # Calcular el total, promedio y máximo para cada tipo de evento
    total_tipo_1 = conteos['Amarilla'].sum()
    promedio_tipo_1 = conteos['Amarilla'].mean()
    maximo_tipo_1 = conteos['Amarilla'].max()

    total_tipo_2_y_3 = conteos['Roja'].sum()
    promedio_tipo_2_y_3 = conteos['Roja'].mean()
    maximo_tipo_2_y_3 = conteos['Roja'].max()

    # Definir la posición de las barras
    x = np.arange(len(conteos))  # Las posiciones de las horas (de 7 a 7+24)
    width = 0.35  # Ancho de las barras

    # Crear el subplot
    fig, ax = plt.subplots(figsize=(15, 8))

    # Plot de las dos series de datos
    bars1 = ax.bar(x - width/2, conteos['Amarilla'], width, label='Amarilla', color='yellow')
    bars2 = ax.bar(x + width/2, conteos['Roja'], width, label='Roja', color='red')

    # Etiquetas y título
    ax.set_xlabel('Horas del día')
    ax.set_ylabel('Eventos')
    report_date = df.iloc[1]['Start'].strftime('%d/%m/%Y')
    ax.set_title(f'Frecuencia de descargas eléctricas por hora del día {report_date}\nSensor 2 Mirador Evelyn', fontsize=16, pad=20)
    ax.set_xticks(x)
    ax.set_xticklabels([f'{(h % 24):02d}:00' for h in range(7, 24)] + [f'{(h % 24):02d}:00' for h in range(0, 7)])

    # Rotar las etiquetas del eje X
    plt.xticks(rotation=90)

    # Mostrar la cantidad de eventos encima de cada barra
    for bar in bars1:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width() / 2, height, str(int(height)), ha='center', va='bottom', fontsize=10)

    for bar in bars2:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width() / 2, height, str(int(height)), ha='center', va='bottom', fontsize=10)

    # Leyenda
    ax.legend(title='Tipo de Evento', loc='upper right', fontsize=9, title_fontsize=11, shadow=True, fancybox=True, facecolor='silver', edgecolor='black')

    # Agregar la grilla
    ax.grid(True, which='both', axis='both', linestyle='--', color='gray', alpha=0.5)

    # Crear la tabla con métricas
    table_data = [
        [total_tipo_1, total_tipo_2_y_3],
        [round(conteos['Amarilla'].mean()), round(conteos['Roja'].mean())],
        [maximo_tipo_1, maximo_tipo_2_y_3]
    ]
    row_labels = ['Total', 'Promedio', 'Máximo']
    column_labels = ['Alerta Amarilla', 'Alerta Roja']

    # Añadir la tabla a la gráfica
    table = ax.table(cellText=table_data, rowLabels=row_labels, colLabels=column_labels, loc='bottom', cellLoc='center', colLoc='center', bbox=[0.1, -0.533, 0.8, 0.3])

    # Personalizar las celdas de la tabla
    for (i, j), cell in table.get_celld().items():
        if j == -1:  # Títulos de las filas
            cell.set_fontsize(10)
            cell.set_text_props(weight='bold')
            cell.set_facecolor('#ffcccb')
            cell.set_text_props(color='black')
        if i == 0:  # Títulos de las columnas
            cell.set_fontsize(10)
            cell.set_text_props(weight='bold')
            cell.set_facecolor('#4CAF50')
            cell.set_text_props(color='white')

    plt.tight_layout()

    # Devolver el gráfico sin mostrarlo
    return ax


# ************************ UTILS
# Función para formatear el tiempo como horas:minutos
def format_duration(td):
    if isinstance(td, pd.Timedelta):
        total_seconds = td.total_seconds()
        hours = int(total_seconds // 3600)
        minutes = int((total_seconds % 3600) // 60)
        return f"{hours:02d}:{minutes:02d}"
    return "00:00"  # En caso de que el valor no sea un Timedelta


def add_row(df, row_data, position=1):
    # Crear un DataFrame con la nueva fila
    new_row = pd.DataFrame([row_data])
    
    # Añadir la fila al inicio o al final
    if position == 1:
        df = pd.concat([new_row, df], ignore_index=True)
    elif position == -1:
        df = pd.concat([df, new_row], ignore_index=True)
    else:
        raise ValueError("El parámetro 'position' debe ser 1 (inicio) o -1 (final).")
    
    return df


def organize_data(df):
    # Verificar las descripciones únicas en la columna 'Description'
    # Las descrpciones pueden ser: Alarm, Alarm Flash, Caution, Caution flash,
    # etc.
    # Ignorar valores nulos - El txt se descarga del sistema con valores nulos
    #unique_descriptions = df['Description'].dropna().unique()

    # Para el reporte diario solamente las descripciones Alarm y Caution son de
    # interés
    data = df[df['Description'].isin(['Alarm', 'Caution'])]

    # Obtener los registros para el Type Start y End
    data_start = data[['Start', 'Description']].copy()
    data_start.rename(columns={'Start': 'Date'}, inplace=True)
    data_start['Type'] = 'Start'
    data_end = data[['End', 'Description']].copy()
    data_end.rename(columns={'End': 'Date'}, inplace=True)
    data_end['Type'] = 'End'

    # Combinar los DataFrames
    combined_data = pd.concat([data_start, data_end], ignore_index=True)

    # Ordenar por la columna 'Date'
    combined_data['Date'] = pd.to_datetime(combined_data['Date'])  # Asegurar formato de fecha
    combined_data = combined_data.sort_values(by='Date')

    # Reiniciar el índice después de ordenar
    combined_data.reset_index(drop=True, inplace=True)

    # Obtener la primera fecha del DataFrame
    today = combined_data['Date'].iloc[0]

    # Asegurarse de que 'today' sea un objeto datetime
    if not isinstance(today, pd.Timestamp):
        today = pd.to_datetime(today)

    # Verificar si la hora es antes de las 7:00 AM
    #if today.time() > pd.Timestamp("09:00:00").time():
    #    # Crear start_day con la fecha del día actual a las 00:00 horas
    #    start_day = pd.Timestamp(today.date())  # Esto pone la hora en 00:00 automáticamente
    #    end_day = start_day + pd.Timedelta(days=1)
    #else:
    #    # Crear start_day con la fecha del día actual a las 07:00 horas
    #    start_day = pd.Timestamp(today.date()) + pd.Timedelta(hours=7)
    #    end_day = start_day + pd.Timedelta(days=1)

    start_day = pd.Timestamp(today.date()) + pd.Timedelta(hours=7)
    end_day = start_day + pd.Timedelta(days=1)

    combined_data['Status'] = None
    combined_data['Duration'] = None
    
    # Agregar una nueva fila al inicio
    organized_data = add_row(
        combined_data,
        row_data={'Date': start_day, 'Description': None, 'Type': None, 'Status': 'White', 'Duration': None},
        position=1
    )

    organized_data = add_row(
        organized_data,
        row_data={'Date': end_day, 'Description': None, 'Type': None, 'Status': 'White', 'Duration': None},
        position=-1
    )

    # Calcular la duración entre cada par de filas consecutivas
    for i in range(1, len(organized_data)):
        # Restar el Date de la fila anterior del Date de la fila actual
        duration = organized_data.loc[i, 'Date'] - organized_data.loc[i - 1, 'Date']
        organized_data.loc[i-1, 'Duration'] = duration  # Convertir a segundos

    return organized_data


def set_status(organized_data):
    is_alarm_before = False
    is_end_after = False

    # Llenar la columna 'Status' con base en las reglas
    for i in range(len(organized_data)):
        # ******************************* ALARM ********************************
        if organized_data.loc[i, 'Description'] == 'Alarm':
            # **************************************************** ALARM - START
            if organized_data.loc[i, 'Type'] == 'Start':
                organized_data.loc[i, 'Status'] = 'Red'


            # ****************************************************** ALARM - END
            elif organized_data.loc[i, 'Type'] == 'End':
                try:
                  if organized_data.loc[i+1, 'Description'] == 'Alarm' and organized_data.loc[i+1, 'Type'] == 'End':
                      is_end_after = True
                except:
                  pass

                if is_end_after:
                    organized_data.loc[i, 'Status'] = 'Red'
                    is_end_after = False
                else:
                    organized_data.loc[i, 'Status'] = 'Yellow'



        # ************************* CAUTION ************************************
        elif organized_data.loc[i, 'Description'] == 'Caution':

            # **************************************************** CAUTION - END
            if organized_data.loc[i, 'Type'] == 'End':
                duration_hours = organized_data.loc[i, 'Duration'].seconds / 3600  # Calcular la duración en horas
                if duration_hours > 1:
                    organized_data.loc[i, 'Status'] = 'White'
                else:
                    organized_data.loc[i, 'Status'] = 'Gray'

            # ************************************************** CAUTION - START
            elif organized_data.loc[i, 'Type'] == 'Start':
                try:
                  if organized_data.loc[i-1, 'Description'] == 'Alarm':
                      is_alarm_before = True
                except:
                  pass

                if is_alarm_before:
                    organized_data.loc[i, 'Status'] = 'Red'
                    is_alarm_before = False
                else:
                    organized_data.loc[i, 'Status'] = 'Yellow'

    final_data = organized_data.copy()

    # Validar si existen "rangos de Status" para sumarlos
    new_data = []
    status = None

    for i in range(len(final_data)):
      if final_data.loc[i, 'Description'] == None:
          new_data.append(final_data.loc[i].copy())
      else:
          if status != final_data.loc[i, 'Status']:
            status = final_data.loc[i, 'Status']
            new_data.append(final_data.loc[i].copy())
          else:
            new_data[-1]['Duration'] += final_data.loc[i, 'Duration']
    new_data = pd.DataFrame(new_data)

    return new_data


def get_daily_plot(final_data):
    # Variables para las duraciones totales por tipo
    total_yellow_duration = pd.Timedelta(0)
    total_gray_duration = pd.Timedelta(0)
    total_red_duration = pd.Timedelta(0)
    total_white_duration = pd.Timedelta(0)

    # Sumar las duraciones de cada tipo
    for _, row in final_data.iterrows():
        if isinstance(row['Duration'], pd.Timedelta):  # Ignorar duraciones vacías o None
            if row['Status'] == 'Yellow':
                total_yellow_duration += row['Duration']
            elif row['Status'] == 'Gray':
                total_gray_duration += row['Duration']
            elif row['Status'] == 'Red':
                total_red_duration += row['Duration']
            elif row['Status'] == 'White':
                total_white_duration += row['Duration']

    status_colors = {
        'Gray': 'grey',
        'White': 'white',
        'Yellow': 'yellow',
        'Red': 'red'
    }

    fig, ax = plt.subplots(figsize=(15, 8))

    start_date = final_data['Date'].min()
    end_date = final_data['Date'].max()
    ax.set_xlim(start_date, end_date)
    ax.xaxis.set_major_locator(mdates.HourLocator(interval=2))
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%H:%M'))


    # Plot the bars
    for i, row in final_data.iterrows():
        # Condición para omitir duraciones de cero
        if isinstance(row['Duration'], pd.Timedelta) and row['Duration'] > pd.Timedelta(0):
            start = row['Date']
            duration_in_minutes = row['Duration'].total_seconds() / 60
            color = status_colors.get(row['Status'], 'black')
            ax.barh(0, width=duration_in_minutes, left=start, color=color, edgecolor='none')
            
            # Formato de duración como hh:mm
            total_seconds = row['Duration'].total_seconds()
            hours = int(total_seconds // 3600)
            minutes = int((total_seconds % 3600) // 60)
            duration_text = f"{hours:02d}:{minutes:02d}"
            
            # Añadir texto verticalmente centrado o hacia abajo para barras blancas
            if color in ['white']:  # Verifica si la barra es blanca
                ax.text(start + row['Duration'] / 2, -0.3, duration_text, ha='center',
                        fontsize=9, color='black', rotation=90)
            else:
                ax.text(start + row['Duration'] / 2, 0, duration_text, ha='center',
                        fontsize=9, color='black', rotation=90)

    ax.set_yticks([])

    # Obtener la fecha DD/MM/YYYY de la segunda columna de date
    report_date = final_data.iloc[1]['Date'].strftime('%d/%m/%Y')
    ax.set_title(f'{report_date} - Sensores Mina', fontsize=16, pad=20, loc='left')

    # Remove x-axis label
    ax.set_xlabel('')

    # Set y-axis label
    ax.set_ylabel('Tipo de Alerta')
        # Filtrar solo las alertas de interés para la leyenda
    legend_status = ['Red', 'Yellow', 'Grey']
    legend_patches = [
        mpatches.Patch(color=color, label='Alerta Roja' if status == 'Red' else 
                                        'Alerta Amarilla' if status == 'Yellow' else 
                                        'Libre entre Alertas <=1hr' if status == 'Grey' else status) 
        for status, color in status_colors.items() if status in legend_status
    ]

    # Crear la leyenda arriba a la derecha, encima de la gráfica
    ax.legend(handles=legend_patches, loc='lower right', bbox_to_anchor=(1, 1.05), ncol=5)

    # Crear la tabla con las métricas
    total_alert_duration = total_yellow_duration + total_red_duration
    total_combined_duration = total_alert_duration + total_gray_duration


    table_data = [
        ['Tiempo Alerta Amarilla', format_duration(total_yellow_duration)],
        ['Tiempo Alerta Roja', format_duration(total_red_duration)],
        ['Total Alertas (Amarilla+Roja)', format_duration(total_alert_duration)],
        ['Total Tiempo Libre Entre Alertas (<=1Hr)', format_duration(total_gray_duration)],
        ['Total Alertas + Tiempo Libre Entre Alertas', format_duration(total_combined_duration)]
    ]


    # Añadir la tabla a la gráfica
    table = ax.table(cellText=table_data, loc='bottom', cellLoc='center', colLoc='center', bbox=[0.1, -0.533, 0.8, 0.3])

    # Personalizar las celdas de la tabla
    for (i, j), cell in table.get_celld().items():
        if j == -1:  # Títulos de las filas
            cell.set_fontsize(10)
            cell.set_text_props(weight='bold')
            cell.set_facecolor('#ffcccb')
            cell.set_text_props(color='black')
        if i == 0:  # Títulos de las columnas
            cell.set_fontsize(10)
            cell.set_text_props(weight='bold')
            cell.set_facecolor('#4CAF50')
            cell.set_text_props(color='white')

    plt.tight_layout()

    return ax


def generate_reports(df):
    report_date_start = pd.to_datetime(df.iloc[1]['Start']).strftime('%d/%m/%Y')
    report_date_end = pd.to_datetime(df.iloc[-1]['Start']).strftime('%d/%m/%Y')
        
    organized_data = organize_data(df)
    final_data = set_status(organized_data)

    ax_1 = get_daily_plot(final_data)
    ax_2 = plot_eventos(df)
    

    # Mostrar la primera gráfica en Streamlit
    st.subheader("Gráfica de Alertas por Día")
    st.pyplot(ax_1.figure)

    # Mostrar la segunda gráfica en Streamlit
    st.subheader("Frecuencia de Alertas por Hora")
    st.pyplot(ax_2.figure)


    # ******************************************

    img_buf_1 = BytesIO()
    ax_1.figure.savefig(img_buf_1, format='png')
    img_buf_1.seek(0)

    img_buf_2 = BytesIO()
    ax_2.figure.savefig(img_buf_2, format='png')
    img_buf_2.seek(0)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    header = doc.sections[0].header
    header_table = header.add_table(rows=1, cols=2, width=doc.sections[0].page_width)
    header_table.columns[0].width = Pt(50)

    cell_logo = header_table.cell(0, 0)
    cell_logo.paragraphs[0].add_run().add_picture("images/logo_doc.PNG", width=Pt(100))

    cell_title = header_table.cell(0, 1)
    header_table.cell(0, 1).width = Pt(1250)
    title_paragraph = cell_title.paragraphs[0]
    title_paragraph.add_run("REPORTE DIARIO DE ALERTAS POR DESCARGAS ELÉCTRICAS ATMOSFÉRICAS\n").bold = True
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle_paragraph = cell_title.add_paragraph(
        f"De: {report_date_start} 07:00 horas\tA: {report_date_end} 07:00 horas"
    )
    subtitle_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #doc.add_paragraph("Gráfica de alertas diarias:")
    doc.add_picture(img_buf_1, width=Pt(650))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #doc.add_paragraph("Gráfica de eventos por hora:")
    doc.add_picture(img_buf_2, width=Pt(650))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    # Obtener la fecha actual
    now = datetime.now()

    # Formatear la fecha manualmente
    today_date = f"{now.day} de {meses_es[now.month]} del {now.year}"

    # Configurar el pie de página
    footer = section.footer  # Acceder al footer de la sección
    footer_paragraph = footer.paragraphs[0]  # Crear un párrafo dentro del footer

    # Configurar el contenido del footer
    footer_paragraph.text = (
        f"{today_date} Supervisión de Mantenimiento en Telecomunicaciones"
    )
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar el texto

    # Ajustar el formato del texto del footer
    for run in footer_paragraph.runs:
        run.font.size = Pt(10)  # Ajustar el tamaño de la fuente

    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)

    st.download_button(
        label="Descargar Informe",
        data=doc_buffer,
        file_name=f"informe_generado-{report_date_start}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    

# Función para cargar y mostrar archivo
def cargar_archivo():
    archivo = st.file_uploader("Selecciona un archivo de texto (.txt)", type=["txt"])
    
    if archivo is not None:
        try:
            # Leer el archivo como DataFrame
            df = pd.read_csv(archivo, sep='\t')
            st.success(f"Archivo cargado correctamente: {archivo.name}")
            
            # Mostrar los datos en un DataFrame interactivo
            st.dataframe(df)

            
            # Botón para generar reportes (aunque no está implementado, lo mostramos)
            if st.button("Generar Reportes"):
                st.write("Generando reportes...")
                generate_reports(df)
        
        except Exception as e:
            st.error(f"Ocurrió un error al leer el archivo: {e}")
    else:
        st.warning("Por favor, selecciona un archivo de texto (.txt)")

# Mostrar el logo (asegúrate de tener el archivo en la ruta correcta)
try:
    imagen = Image.open("images/logo.png")
    imagen = imagen.resize((100, 50), Image.Resampling.LANCZOS)
    st.image(imagen, use_container_width=False)
except FileNotFoundError:
    st.warning("No se encontró la imagen en la ruta especificada.")


# Título largo de la app
st.title("Generador de Reportes Diarios de Alertas por Descargas Eléctricas Atmosféricas")
# Llamar a la función para cargar el archivo
cargar_archivo()
